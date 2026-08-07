#!/usr/bin/env python3
"""Apply stable journal-style layout to the Pandoc-generated manuscript DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, twips: int = 55) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row, repeat_header: bool = False) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        properties.append(OxmlElement("w:tblHeader"))


def set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "B7B7B7")


def format_table(table, widths: list[float]) -> None:
    table.style = "TableNormal"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "pct")
    table_width.set(qn("w:w"), "5000")

    for column, width in zip(table.columns, widths, strict=True):
        column.width = Inches(width)
        for cell in column.cells:
            cell.width = Inches(width)

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row, repeat_header=row_index == 0)
        for cell_index, cell in enumerate(row.cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, "1F4E78")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "EAF2F8")
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal"
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if cell_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    run.font.bold = row_index == 0
                    if row_index == 0:
                        run.font.color.rgb = RGBColor(255, 255, 255)


def format_paragraphs(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(15, 76, 101)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^Table \d+\.", text):
            paragraph.style = "Normal"
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
                run.font.bold = True
        elif re.match(r"^Figure \d+\.", text):
            paragraph.style = "Normal"
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.input)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)

    format_paragraphs(document)
    for table in document.tables:
        width = 6.9 / max(len(table.columns), 1)
        widths = [width] * len(table.columns)
        format_table(table, widths)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
